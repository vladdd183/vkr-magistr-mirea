#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Mm, Pt


ROOT = Path(__file__).parent
META_PATH = ROOT / "meta.typ"
CONTENT_SRC = ROOT / "content"
CONTENT_DOCX = ROOT / "content-docx"
DOCX_OUTPUT = ROOT / "main.docx"

DOCX_IMAGE_WIDTH_OVERRIDES = {
    "diagrams/output/c4/c4-deployment.png": "70%",
    "diagrams/output/c4/c4-container.png": "105%",
    "diagrams/output/c4/c4-context.png": "105%",
    "diagrams/output/06-sequence.png": "105%",
    "diagrams/output/c4/recommendation-pipeline.png": "105%",
}

DOCX_FIGURES_WITH_FRAME = {"1.1", "1.2", "2.2"}

DOCX_FIGURES_WITH_TB_BORDER = {"2.3", "2.4", "2.5", "2.6", "2.7"}

FIGURE_CAPTION_SUFFIX = " [разработано автором]"

UNNUMBERED_H1 = {
    "Реферат",
    "Содержание",
    "Перечень сокращений и обозначений",
    "Введение",
    "Заключение",
    "Conclusion",
    "Список использованных источников",
    "Приложения",
}

SPECIAL_BOLD_SUBHEADING_MARKER = "[[DOCX-BOLD-SUBHEADING]] "
CODE_BLOCK_MARKER = "[[DOCX-CODE-BLOCK]]"


@dataclass
class Formatting:
    margin_top_mm: float = 20.0
    margin_bottom_mm: float = 20.0
    margin_left_mm: float = 30.0
    margin_right_mm: float = 15.0
    font_size_pt: float = 14.0
    paragraph_indent_cm: float = 1.25
    line_spacing: float = 1.5


@dataclass
class Meta:
    doc_type: str
    title: str
    discipline: str
    student_name: str
    student_group: str
    supervisor_name: str
    supervisor_title: str
    university: str
    institute: str
    department: str
    city: str
    year: str
    abbreviations: list[tuple[str, str]]
    formatting: Formatting


class DocxBuilder:
    def __init__(self, meta: Meta) -> None:
        self.meta = meta
        self.doc = Document()
        self.figure_counter = 0
        self.table_counter = 0
        self.figure_in_chapter = 0
        self.table_in_chapter = 0
        self.chapter_counter = 0
        self.h2_counter = 0
        self.h3_counter = 0
        self.current_chapter_number: int | None = None
        self.listing_counter = 0
        self.listing_in_chapter = 0
        self.pending_space_before_after_block = False
        self._setup_document()

    @property
    def content_width_cm(self) -> float:
        return 21.0 - (self.meta.formatting.margin_left_mm + self.meta.formatting.margin_right_mm) / 10.0

    def _setup_document(self) -> None:
        section = self.doc.sections[0]
        fmt = self.meta.formatting
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(fmt.margin_top_mm)
        section.bottom_margin = Mm(fmt.margin_bottom_mm)
        section.left_margin = Mm(fmt.margin_left_mm)
        section.right_margin = Mm(fmt.margin_right_mm)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(fmt.font_size_pt)
        normal_par = normal.paragraph_format
        normal_par.first_line_indent = Cm(fmt.paragraph_indent_cm)
        normal_par.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal_par.space_after = Pt(0)
        normal_par.space_before = Pt(0)
        self._set_style_fonts(normal.element)

        self._configure_heading_style("Heading 1", Pt(14), True, first_line_indent=0, space_before=Pt(0), space_after=Pt(0))
        self._configure_heading_style("Heading 2", Pt(14), True, first_line_indent=0, space_before=Pt(0), space_after=Pt(0))
        self._configure_heading_style("Heading 3", Pt(14), True, first_line_indent=0, space_before=Pt(0), space_after=Pt(0))

        self._enable_update_fields_on_open()

    def _configure_heading_style(
        self,
        style_name: str,
        font_size: Pt,
        bold: bool,
        *,
        first_line_indent: float,
        space_before: Pt,
        space_after: Pt,
    ) -> None:
        style = self.doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = font_size
        style.font.bold = bold
        if style.font.color is not None:
            style.font.color.rgb = None
        par = style.paragraph_format
        par.first_line_indent = Cm(first_line_indent) if first_line_indent else None
        par.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        par.space_before = space_before
        par.space_after = space_after
        self._set_style_fonts(style.element)
        self._force_black_text(style.element)

    def _set_style_fonts(self, style_element) -> None:
        r_pr = style_element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            style_element.append(r_pr)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(key), "Times New Roman")

    def _force_black_text(self, style_element) -> None:
        r_pr = style_element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            style_element.append(r_pr)
        color = r_pr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
            if color.get(qn(attr)) is not None:
                del color.attrib[qn(attr)]

    def _set_run_font(self, run, *, size: Pt | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
        run.font.name = "Times New Roman"
        if size is not None:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run._element.insert(0, r_pr)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(key), "Times New Roman")
        color = r_pr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
            if color.get(qn(attr)) is not None:
                del color.attrib[qn(attr)]

    def _add_page_numbers_to_section(self, section) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        begin = paragraph.add_run()
        begin_field = OxmlElement("w:fldChar")
        begin_field.set(qn("w:fldCharType"), "begin")
        begin._element.append(begin_field)

        instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        instr._element.append(instr_text)

        end = paragraph.add_run()
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end._element.append(end_field)

        for run in (begin, instr, end):
            self._set_run_font(run, size=Pt(14))

    def start_numbered_section(self) -> None:
        new_section = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        fmt = self.meta.formatting
        new_section.page_width = Mm(210)
        new_section.page_height = Mm(297)
        new_section.top_margin = Mm(fmt.margin_top_mm)
        new_section.bottom_margin = Mm(fmt.margin_bottom_mm)
        new_section.left_margin = Mm(fmt.margin_left_mm)
        new_section.right_margin = Mm(fmt.margin_right_mm)
        self._add_page_numbers_to_section(new_section)

    def _enable_update_fields_on_open(self) -> None:
        settings = self.doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def add_title_page(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for line in (
            "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ",
            "РОССИЙСКОЙ ФЕДЕРАЦИИ",
            "",
            "Федеральное государственное бюджетное образовательное учреждение",
            "высшего образования",
            "",
            f"«{self.meta.university}»",
            "",
            self.meta.institute,
            self.meta.department,
        ):
            run = p.add_run(f"{line}\n")
            self._set_run_font(run, size=Pt(14), bold=("«" in line or line == self.meta.department))

        self._add_spacer()
        self._add_spacer()

        middle = self.doc.add_paragraph()
        middle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        middle.paragraph_format.first_line_indent = None
        middle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self._add_formatted_text(middle, f"{doc_type_name(self.meta.doc_type)}\n", bold=True)
        self._add_formatted_text(middle, f'по дисциплине: «{self.meta.discipline}»\n')
        self._add_formatted_text(middle, "Тема:\n")
        self._add_formatted_text(middle, f"«{self.meta.title}»", bold=True)

        self._add_spacer()
        self._add_spacer()

        author = self.doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.LEFT
        author.paragraph_format.first_line_indent = None
        author.paragraph_format.left_indent = Cm(8.5)
        author.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self._add_formatted_text(author, f"Выполнил: студент группы {self.meta.student_group}\n")
        self._add_formatted_text(author, f"{self.meta.student_name}\n\n")
        self._add_formatted_text(author, f"Руководитель: {self.meta.supervisor_title}\n")
        self._add_formatted_text(author, self.meta.supervisor_name)

        self._add_spacer()
        self._add_spacer()

        bottom = self.doc.add_paragraph()
        bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom.paragraph_format.first_line_indent = None
        self._add_formatted_text(bottom, f"{self.meta.city}\n{self.meta.year}")

        self.doc.add_page_break()

    def _add_spacer(self) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def _add_formatted_text(self, paragraph, text: str, *, bold: bool = False, italic: bool = False, size: Pt | None = None) -> None:
        run = paragraph.add_run(self._normalize_text(text))
        self._set_run_font(run, size=size or Pt(self.meta.formatting.font_size_pt), bold=bold, italic=italic)

    def _add_field(self, paragraph, instruction: str) -> None:
        begin = paragraph.add_run()
        begin_field = OxmlElement("w:fldChar")
        begin_field.set(qn("w:fldCharType"), "begin")
        begin._element.append(begin_field)
        self._set_run_font(begin, size=Pt(self.meta.formatting.font_size_pt))

        instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {instruction} "
        instr._element.append(instr_text)
        self._set_run_font(instr, size=Pt(self.meta.formatting.font_size_pt))

        separate = paragraph.add_run()
        separate_field = OxmlElement("w:fldChar")
        separate_field.set(qn("w:fldCharType"), "separate")
        separate._element.append(separate_field)
        self._set_run_font(separate, size=Pt(self.meta.formatting.font_size_pt))

        # Give Word a visible placeholder so the field is not blank
        # before the document updates fields on open.
        placeholder = paragraph.add_run("1")
        self._set_run_font(placeholder, size=Pt(self.meta.formatting.font_size_pt))

        end = paragraph.add_run()
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end._element.append(end_field)
        self._set_run_font(end, size=Pt(self.meta.formatting.font_size_pt))

    def _add_centered_title(self, text: str, *, bold: bool = True) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        self._add_formatted_text(paragraph, text.upper(), bold=bold)

    def add_toc(self) -> None:
        self._add_centered_title("Содержание")

        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = None
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        begin = paragraph.add_run()
        begin_field = OxmlElement("w:fldChar")
        begin_field.set(qn("w:fldCharType"), "begin")
        begin._element.append(begin_field)

        instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        instr._element.append(instr_text)

        separate = paragraph.add_run()
        separate_field = OxmlElement("w:fldChar")
        separate_field.set(qn("w:fldCharType"), "separate")
        separate._element.append(separate_field)

        placeholder = paragraph.add_run("Оглавление будет обновлено при открытии документа.")
        self._set_run_font(placeholder, size=Pt(12), italic=True)

        end = paragraph.add_run()
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end._element.append(end_field)

    def add_abbreviations(self) -> None:
        if not self.meta.abbreviations:
            return

        self.add_page_break()
        self._add_centered_title("Перечень сокращений и обозначений")

        abbreviations = sort_abbreviations(self.meta.abbreviations)
        table = self.doc.add_table(rows=len(abbreviations), cols=3)
        self._clear_table_borders(table)

        for idx, (abbr, full) in enumerate(abbreviations):
            self._set_cell(
                table.rows[idx].cells[0],
                abbr,
                size=Pt(self.meta.formatting.font_size_pt),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            self._set_cell(
                table.rows[idx].cells[1],
                "—",
                size=Pt(self.meta.formatting.font_size_pt),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            self._set_cell(
                table.rows[idx].cells[2],
                full,
                size=Pt(self.meta.formatting.font_size_pt),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )

        first_width = self._estimate_text_width_cm([abbr for abbr, _ in abbreviations], min_cm=1.6, max_cm=4.5)
        second_width = self._estimate_text_width_cm(["—"], min_cm=0.45, max_cm=0.8)
        third_width = max(6.0, self.content_width_cm - first_width - second_width - 0.6)
        self._set_column_width(table, 0, first_width)
        self._set_column_width(table, 1, second_width)
        self._set_column_width(table, 2, third_width)

    def _set_table_borders(self, table) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = parse_xml(
            '<w:tblBorders %s>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>' % nsdecls("w")
        )
        tbl_pr.append(borders)

    def _clear_table_borders(self, table) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = parse_xml(
            '<w:tblBorders %s>'
            '  <w:top w:val="nil"/>'
            '  <w:left w:val="nil"/>'
            '  <w:bottom w:val="nil"/>'
            '  <w:right w:val="nil"/>'
            '  <w:insideH w:val="nil"/>'
            '  <w:insideV w:val="nil"/>'
            '</w:tblBorders>' % nsdecls("w")
        )
        tbl_pr.append(borders)

    def _set_column_width(self, table, column_idx: int, width_cm: float) -> None:
        table.autofit = False
        table.columns[column_idx].width = Cm(width_cm)
        for row in table.rows:
            row.cells[column_idx].width = Cm(width_cm)

    def _estimate_text_width_cm(self, texts: list[str], *, min_cm: float, max_cm: float) -> float:
        longest = 0
        longest_word = 0
        for text in texts:
            normalized = self._normalize_text(text)
            for line in normalized.splitlines():
                stripped = line.strip()
                longest = max(longest, len(stripped))
                for word in re.split(r"[\s,;:.()\\[\\]«»\"/]+", stripped):
                    if word:
                        longest_word = max(longest_word, len(word))
        estimate = max(0.18 * longest + 0.6, 0.22 * longest_word + 0.8)
        return max(min_cm, min(max_cm, estimate))

    def _set_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        size: Pt | None = None,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        vertical_align=WD_ALIGN_VERTICAL.CENTER,
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(self._normalize_table_cell_text(text))
        self._set_run_font(run, size=size or Pt(12), bold=bold)
        cell.vertical_alignment = vertical_align

    def _normalize_table_cell_text(self, text: str) -> str:
        compact = re.sub(r"\s+", " ", self._normalize_text(text)).strip()
        if not compact:
            return ""

        if compact.startswith("— "):
            bullet_items = [item.strip(" ;") for item in re.split(r"(?:^|\s+)—\s+", compact) if item.strip(" ;")]
            if len(bullet_items) > 1:
                return " ".join(bullet_items)

        if re.match(r"^\d+\.\s+", compact):
            numbered_items = [item.strip(" ;") for item in re.split(r"(?:^|\s+)\d+\.\s+", compact) if item.strip(" ;")]
            if len(numbered_items) > 1:
                return " ".join(numbered_items)

        return compact

    def _normalize_text(self, text: str) -> str:
        return text.replace("→", "—")

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    def add_abstract_file(self, path: Path) -> None:
        lines = [line.strip() for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
        if not lines:
            return

        heading = lines.pop(0)
        if heading.startswith("# "):
            self._add_heading(heading[2:].strip(), 1, break_before_h1=False)

        if lines:
            stats_par = self.doc.add_paragraph()
            stats_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats_par.paragraph_format.first_line_indent = None
            stats_par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            self._add_inline_runs(stats_par, lines.pop(0))

        if lines:
            keywords = lines.pop(0)
            key_par = self.doc.add_paragraph()
            key_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            key_par.paragraph_format.first_line_indent = None
            self._add_inline_runs(key_par, keywords)
        for line in lines:
            self._add_paragraph(line)

    def add_markdown_file(self, path: Path, *, break_before_h1: bool = True) -> None:
        lines = path.read_text(encoding="utf-8").splitlines()
        i = 0
        while i < len(lines):
            raw = lines[i].rstrip()
            stripped = raw.strip()

            if not stripped:
                i += 1
                continue

            if stripped.startswith("Таблица:"):
                i += 1
                continue

            if stripped == "<!-- pagebreak -->":
                self.add_page_break()
                i += 1
                continue

            if is_markdown_table(lines, i):
                i = self._add_markdown_table(lines, i)
                continue

            image = parse_image_line(stripped)
            if image:
                self._add_image(path, image["caption"], image["target"], image["width"])
                i += 1
                continue

            if stripped.startswith(SPECIAL_BOLD_SUBHEADING_MARKER):
                self._add_special_bold_subheading(stripped[len(SPECIAL_BOLD_SUBHEADING_MARKER) :].strip())
                i += 1
                continue

            if stripped.startswith(CODE_BLOCK_MARKER):
                caption = stripped[len(CODE_BLOCK_MARKER):].strip()
                i += 1
                code_lines: list[str] = []
                while i < len(lines):
                    code_line = lines[i]
                    if code_line.strip() == "[[/DOCX-CODE-BLOCK]]":
                        i += 1
                        break
                    code_lines.append(code_line)
                    i += 1
                self._add_code_block(caption, "\n".join(code_lines))
                continue

            heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
            if heading:
                level = len(heading.group(1))
                title = heading.group(2).strip()
                self._add_heading(title, level, break_before_h1=break_before_h1)
                i += 1
                continue

            list_match = re.match(r"^([-*]|\d+\.)\s+(.+)$", stripped)
            if list_match:
                marker = list_match.group(1)
                ordered = marker.endswith(".") and marker[0].isdigit()
                items: list[str] = []
                while i < len(lines):
                    candidate = lines[i].strip()
                    if not candidate:
                        probe = i + 1
                        while probe < len(lines) and not lines[probe].strip():
                            probe += 1
                        if ordered and probe < len(lines) and re.match(r"^\d+\.\s+(.+)$", lines[probe].strip()):
                            i = probe
                            candidate = lines[i].strip()
                        else:
                            break
                    candidate_match = re.match(r"^([-*]|\d+\.)\s+(.+)$", candidate)
                    if not candidate_match:
                        break
                    items.append(candidate_match.group(2).strip())
                    i += 1
                self._add_list(items, ordered=ordered)
                continue

            paragraph_lines = [stripped]
            i += 1
            while i < len(lines):
                candidate = lines[i].strip()
                if (
                    not candidate
                    or candidate.startswith("Таблица:")
                    or candidate == "<!-- pagebreak -->"
                    or candidate.startswith(SPECIAL_BOLD_SUBHEADING_MARKER)
                    or re.match(r"^(#{1,3})\s+", candidate)
                    or parse_image_line(candidate)
                    or is_markdown_table(lines, i)
                    or re.match(r"^([-*]|\d+\.)\s+(.+)$", candidate)
                ):
                    break
                paragraph_lines.append(candidate)
                i += 1
            self._add_paragraph(" ".join(paragraph_lines))

    def _add_heading(self, title: str, level: int, *, break_before_h1: bool) -> None:
        if level == 1 and self.doc.paragraphs and self.doc.paragraphs[-1].text.strip() and break_before_h1:
            self.add_page_break()

        display = title
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        if level == 1:
            if title not in UNNUMBERED_H1:
                self.chapter_counter += 1
                self.current_chapter_number = self.chapter_counter
                self.figure_in_chapter = 0
                self.table_in_chapter = 0
                self.listing_in_chapter = 0
                self.h2_counter = 0
                self.h3_counter = 0
                display = f"{self.chapter_counter} {title}"
            else:
                self.current_chapter_number = None
                self.h2_counter = 0
                self.h3_counter = 0
                display = title.upper()
                alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2 and self.current_chapter_number is not None:
            self.h2_counter += 1
            self.h3_counter = 0
            display = f"{self.current_chapter_number}.{self.h2_counter} {title}"
        elif level == 3 and self.current_chapter_number is not None:
            if self.h2_counter == 0:
                self.h2_counter = 1
            self.h3_counter += 1
            display = f"{self.current_chapter_number}.{self.h2_counter}.{self.h3_counter} {title}"

        paragraph = self.doc.add_paragraph(style=f"Heading {level}")
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.paragraph_format.first_line_indent = Cm(0)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(self.meta.formatting.paragraph_indent_cm)
        if self.pending_space_before_after_block:
            paragraph.paragraph_format.space_before = Pt(6)
            self.pending_space_before_after_block = False
        paragraph.alignment = alignment
        self._add_inline_runs(paragraph, display, bold=True)

    def _add_special_bold_subheading(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(6) if self.pending_space_before_after_block else Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        self.pending_space_before_after_block = False
        self._add_inline_runs_with_size(paragraph, text, bold=True, size=Pt(self.meta.formatting.font_size_pt))

    def _add_paragraph(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(self.meta.formatting.paragraph_indent_cm)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(6) if self.pending_space_before_after_block else Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        self.pending_space_before_after_block = False
        self._add_inline_runs(paragraph, text)

    def _add_list(self, items: list[str], *, ordered: bool) -> None:
        indent_cm = self.meta.formatting.paragraph_indent_cm
        for idx, item in enumerate(items, start=1):
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = Cm(indent_cm)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.space_before = Pt(6) if idx == 1 and self.pending_space_before_after_block else Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            prefix = f"{idx}. " if ordered else "— "
            self._add_inline_runs(paragraph, f"{prefix}{item}")
        self.pending_space_before_after_block = False

    def _add_inline_runs(self, paragraph, text: str, *, bold: bool = False) -> None:
        self._add_inline_runs_with_size(paragraph, text, bold=bold, size=Pt(self.meta.formatting.font_size_pt))

    def _add_inline_runs_with_size(self, paragraph, text: str, *, bold: bool = False, size: Pt | None = None) -> None:
        if "{{NUMPAGES}}" in text:
            parts = text.split("{{NUMPAGES}}")
            for idx, part in enumerate(parts):
                if part:
                    self._add_inline_runs_with_size(paragraph, part, bold=bold, size=size)
                if idx < len(parts) - 1:
                    self._add_field(paragraph, "NUMPAGES")
            return

        plain_text = self._normalize_text(text.replace("**", "").replace("*", ""))
        run = paragraph.add_run(plain_text)
        self._set_run_font(run, size=size or Pt(self.meta.formatting.font_size_pt), bold=bold)

    def _add_image(self, source_path: Path, caption: str, target: str, width_spec: str | None) -> None:
        image_path = resolve_asset_path(source_path, target)
        if not image_path.exists():
            self._add_paragraph(f"[Изображение не найдено: {target}]")
            return

        override = DOCX_IMAGE_WIDTH_OVERRIDES.get(relative_to_root(image_path))
        width_cm = width_from_spec(override or width_spec, self.content_width_cm)
        self.figure_counter += 1
        self.figure_in_chapter += 1
        figure_number = self._chapter_figure_number(self.figure_in_chapter)
        if figure_number in DOCX_FIGURES_WITH_FRAME:
            self._add_bordered_picture(str(image_path), width_cm)
        elif figure_number in DOCX_FIGURES_WITH_TB_BORDER:
            self._add_tb_bordered_picture(str(image_path), width_cm)
        else:
            self.doc.add_picture(str(image_path), width=Cm(width_cm))
        picture_paragraph = self.doc.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.first_line_indent = None

        caption_par = self.doc.add_paragraph()
        caption_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_par.paragraph_format.first_line_indent = None
        caption_par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        caption_par.paragraph_format.space_before = Pt(0)
        caption_par.paragraph_format.space_after = Pt(12)
        prefix = f"Рисунок {figure_number} — "
        normalized_caption = ensure_figure_caption_suffix(caption)
        caption_size = Pt(self.meta.formatting.font_size_pt)
        self._add_formatted_text(caption_par, prefix, bold=False, size=caption_size)
        self._add_inline_runs_with_size(caption_par, normalized_caption, bold=False, size=caption_size)
        self.pending_space_before_after_block = True

    def _add_bordered_picture(self, image_path: str, width_cm: float) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.add_run()
        inline_shape = run.add_picture(image_path, width=Cm(width_cm))
        self._apply_picture_border(inline_shape)

    def _add_tb_bordered_picture(self, image_path: str, width_cm: float) -> None:
        from docx.oxml.ns import qn
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for side in ("top", "bottom"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "8")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tc_borders.append(border)
        for side in ("left", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "nil")
            tc_borders.append(border)
        tc_pr.append(tc_borders)
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "bottom", "left", "right"):
            mar = OxmlElement(f"w:{side}")
            mar.set(qn("w:w"), "0")
            mar.set(qn("w:type"), "dxa")
            tc_mar.append(mar)
        tc_pr.append(tc_mar)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
        tbl_pr = table._tbl.tblPr
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), "0")
        tbl_w.set(qn("w:type"), "auto")
        tbl_pr.append(tbl_w)

    def _apply_picture_border(self, inline_shape) -> None:
        sp_pr_list = inline_shape._inline.xpath(".//pic:spPr")
        if not sp_pr_list:
            return
        sp_pr = sp_pr_list[0]
        for existing in sp_pr.xpath("./a:ln"):
            sp_pr.remove(existing)
        line = OxmlElement("a:ln")
        line.set("w", "12700")
        line.set("cap", "flat")
        line.set("cmpd", "sng")
        line.set("algn", "ctr")

        solid_fill = OxmlElement("a:solidFill")
        color = OxmlElement("a:srgbClr")
        color.set("val", "000000")
        solid_fill.append(color)
        line.append(solid_fill)

        preset_dash = OxmlElement("a:prstDash")
        preset_dash.set("val", "solid")
        line.append(preset_dash)

        round_join = OxmlElement("a:round")
        line.append(round_join)

        sp_pr.append(line)

    def _apply_paragraph_tb_border(self, paragraph) -> None:
        from docx.oxml.ns import qn
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        for side in ("top", "bottom"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "8")
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "000000")
            p_bdr.append(border)
        p_pr.append(p_bdr)

    def _add_code_block(self, caption: str, code: str) -> None:
        self.listing_counter += 1
        self.listing_in_chapter += 1
        listing_number = self._chapter_listing_number(self.listing_in_chapter)

        cap_par = self.doc.add_paragraph()
        cap_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_par.paragraph_format.first_line_indent = Cm(0)
        cap_par.paragraph_format.left_indent = Cm(0)
        cap_par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap_par.paragraph_format.space_before = Pt(6) if self.pending_space_before_after_block else Pt(0)
        cap_par.paragraph_format.space_after = Pt(0)
        cap_size = Pt(self.meta.formatting.font_size_pt)
        self._add_inline_runs_with_size(cap_par, f"Листинг {listing_number} — {caption}", size=cap_size)

        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_borders(table)
        cell = table.rows[0].cells[0]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)

        code_lines = code.rstrip("\n").split("\n")
        for idx, line in enumerate(code_lines):
            display = line if line else " "
            run = paragraph.add_run(display)
            self._set_code_run_font(run, size=Pt(12))
            if idx < len(code_lines) - 1:
                run.add_break()

        self.pending_space_before_after_block = True

    def _set_code_run_font(self, run, *, size: Pt | None = None) -> None:
        font_name = "Courier New"
        run.font.name = font_name
        if size is not None:
            run.font.size = size
        run.font.bold = False
        run.font.italic = False
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run._element.insert(0, r_pr)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(key), font_name)
        color = r_pr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)
        color.set(qn("w:val"), "000000")

    def _chapter_listing_number(self, number: int) -> str:
        chapter = self.current_chapter_number or 1
        return f"{chapter}.{number}"

    def _chapter_figure_number(self, number: int) -> str:
        chapter = self.current_chapter_number or 1
        return f"{chapter}.{number}"

    def _chapter_table_number(self, number: int) -> str:
        chapter = self.current_chapter_number or 1
        return f"{chapter}.{number}"

    def _add_table_caption(self, text: str) -> None:
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.left_indent = Cm(0)
        cap.paragraph_format.right_indent = Cm(0)
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(0)
        self._add_inline_runs_with_size(cap, text, size=Pt(self.meta.formatting.font_size_pt))

    def _add_table_body(self, header_cells: list[str], rows: list[list[str]]) -> None:
        table = self.doc.add_table(rows=1 + len(rows), cols=len(header_cells))
        self._set_table_borders(table)
        for idx, cell in enumerate(header_cells):
            self._set_cell(table.rows[0].cells[idx], cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for row_idx, row in enumerate(rows, start=1):
            for col_idx in range(len(header_cells)):
                text = row[col_idx] if col_idx < len(row) else ""
                self._set_cell(table.rows[row_idx].cells[col_idx], text)
        if len(header_cells) == 2:
            self._apply_two_column_layout(table, header_cells, rows)

    def _add_markdown_table(self, lines: list[str], start: int) -> int:
        caption = ""
        probe = start - 1
        while probe >= 0 and not lines[probe].strip():
            probe -= 1
        if probe >= 0:
            previous = lines[probe].strip()
            if previous.startswith("Таблица:"):
                caption = previous.split(":", 1)[1].strip()

        header_cells = split_md_row(lines[start])
        i = start + 2
        rows: list[list[str]] = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_md_row(lines[i]))
            i += 1

        table_number = None
        if caption:
            self.table_counter += 1
            self.table_in_chapter += 1
            table_number = self._chapter_table_number(self.table_in_chapter)

        if table_number == "1.1" and len(rows) >= 4:
            first_part_rows = rows[:-2]
            second_part_rows = rows[-2:]
            self._add_table_caption(f"Таблица {table_number} — {caption}")
            self._add_table_body(header_cells, first_part_rows)
            self._add_table_caption(f"Продолжение таблицы {table_number}")
            self._add_table_body(header_cells, second_part_rows)
        else:
            if caption and table_number is not None:
                self._add_table_caption(f"Таблица {table_number} — {caption}")
            self._add_table_body(header_cells, rows)

        self.pending_space_before_after_block = True
        return i

    def _apply_two_column_layout(self, table, header_cells: list[str], rows: list[list[str]]) -> None:
        first_texts = [header_cells[0]] + [row[0] for row in rows if row]
        first_width = self._estimate_text_width_cm(first_texts, min_cm=3.2, max_cm=6.8)
        total_width = self.content_width_cm
        second_width = max(6.5, total_width - first_width)
        self._set_column_width(table, 0, first_width)
        self._set_column_width(table, 1, second_width)

    def save(self, path: Path) -> None:
        self.doc.save(str(path))


def doc_type_name(doc_type: str) -> str:
    mapping = {
        "coursework": "КУРСОВАЯ РАБОТА",
        "practice-report": "ОТЧЁТ ПО ПРЕДДИПЛОМНОЙ ПРАКТИКЕ",
        "lab-report": "ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ",
    }
    return mapping.get(doc_type, doc_type.upper())


def parse_meta(path: Path) -> Meta:
    text = path.read_text(encoding="utf-8")
    student_block = extract_tuple_block(text, "#let student = (")
    supervisor_block = extract_tuple_block(text, "#let supervisor = (")
    abbreviations_block = extract_tuple_block(text, "#let abbreviations = (")
    formatting_block = extract_tuple_block(text, "#let formatting = (")

    return Meta(
        doc_type=extract_string_assignment(text, "doc-type"),
        title=extract_string_assignment(text, "title"),
        discipline=extract_string_assignment(text, "discipline"),
        student_name=extract_string_field(student_block, "name"),
        student_group=extract_string_field(student_block, "group"),
        supervisor_name=extract_string_field(supervisor_block, "name"),
        supervisor_title=extract_string_field(supervisor_block, "title"),
        university=extract_string_assignment(text, "university"),
        institute=extract_string_assignment(text, "institute"),
        department=extract_string_assignment(text, "department"),
        city=extract_string_assignment(text, "city"),
        year=extract_string_assignment(text, "year"),
        abbreviations=parse_abbreviations(abbreviations_block),
        formatting=parse_formatting(formatting_block),
    )


def extract_tuple_block(text: str, marker: str) -> str:
    start = text.index(marker) + len(marker) - 1
    return extract_balanced(text, start, "(", ")")


def extract_string_assignment(text: str, name: str) -> str:
    match = re.search(rf'#let\s+{re.escape(name)}\s*=\s*"([^"]+)"', text)
    return match.group(1).strip() if match else ""


def extract_string_field(block: str, name: str) -> str:
    match = re.search(rf"{re.escape(name)}:\s*\"([^\"]+)\"", block)
    return match.group(1).strip() if match else ""


def parse_abbreviations(block: str) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    for match in re.finditer(r'\(abbr:\s*"([^"]+)",\s*full:\s*\[([^\]]+)\]\)', block):
        pairs.append((match.group(1).strip(), cleanup_inline_typst(match.group(2).strip())))
    return pairs


def sort_abbreviations(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    def item_key(item: tuple[str, str]) -> tuple[int, str]:
        abbr = item[0].strip()
        if re.search(r"[А-Яа-яЁё]", abbr):
            group = 0
        elif re.search(r"[A-Za-z]", abbr):
            group = 1
        else:
            group = 2
        normalized = abbr.casefold().replace("ё", "е")
        return (group, normalized)

    return sorted(items, key=item_key)


def parse_formatting(block: str) -> Formatting:
    return Formatting(
        margin_top_mm=extract_mm_value(block, "margin-top", 20.0),
        margin_bottom_mm=extract_mm_value(block, "margin-bottom", 20.0),
        margin_left_mm=extract_mm_value(block, "margin-left", 30.0),
        margin_right_mm=extract_mm_value(block, "margin-right", 15.0),
        font_size_pt=extract_pt_value(block, "font-size", 14.0),
        paragraph_indent_cm=extract_cm_value(block, "paragraph-indent", 1.25),
        line_spacing=extract_em_value(block, "line-spacing", 1.5),
    )


def extract_mm_value(text: str, name: str, default: float) -> float:
    match = re.search(rf"{re.escape(name)}:\s*([\d.]+)mm", text)
    return float(match.group(1)) if match else default


def extract_cm_value(text: str, name: str, default: float) -> float:
    match = re.search(rf"{re.escape(name)}:\s*([\d.]+)cm", text)
    return float(match.group(1)) if match else default


def extract_pt_value(text: str, name: str, default: float) -> float:
    match = re.search(rf"{re.escape(name)}:\s*([\d.]+)pt", text)
    return float(match.group(1)) if match else default


def extract_em_value(text: str, name: str, default: float) -> float:
    match = re.search(rf"{re.escape(name)}:\s*([\d.]+)em", text)
    return float(match.group(1)) if match else default


def extract_balanced(text: str, start: int, open_char: str, close_char: str) -> str:
    depth = 0
    begin = None
    for index in range(start, len(text)):
        char = text[index]
        if char == open_char:
            if begin is None:
                begin = index
            depth += 1
        elif char == close_char and begin is not None:
            depth -= 1
            if depth == 0:
                return text[begin:index + 1]
    raise ValueError(f"Не удалось извлечь блок {open_char}{close_char}")


def replace_code_blocks(text: str) -> str:
    """Convert #code-block("caption", lang: "x")[```...```] into marker-delimited blocks."""
    result: list[str] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        cb_match = re.match(r'#code-block\("([^"]+)"(?:,\s*lang:\s*"[^"]*")?\)\[', stripped)
        if cb_match:
            caption = cb_match.group(1)
            result.append(f"{CODE_BLOCK_MARKER}{caption}")
            i += 1
            while i < len(lines):
                inner = lines[i]
                inner_stripped = inner.strip()
                if inner_stripped.startswith("```"):
                    i += 1
                    continue
                if re.match(r"^\]\s*(<[^>]+>)?$", inner_stripped):
                    i += 1
                    break
                result.append(inner.rstrip())
                i += 1
            result.append("[[/DOCX-CODE-BLOCK]]")
            continue
        result.append(lines[i])
        i += 1
    return "\n".join(result)


def cleanup_table_line(text: str) -> str:
    """Clean up Typst artifacts in table caption and cell lines."""
    stripped_chars = set(text.replace("|", "").replace(" ", ""))
    if stripped_chars <= {"-"}:
        return text
    text = text.replace("~", " ")
    text = text.replace("---", "—")
    text = re.sub(r"(?<!-)--(?!-)", "—", text)
    return text


def cleanup_inline_typst(text: str) -> str:
    text = re.sub(r"#cite-range\((\d+),\s*(\d+)\)", r"[\1-\2]", text)
    text = re.sub(r"#cite-src\(([^)]+)\)", lambda m: "[" + cleanup_citation_numbers(m.group(1)) + "]", text)
    text = re.sub(r"#context\s+", "", text)
    text = re.sub(r"(?:#)?counter\([^)]+\)\.\w+\([^)]*\)(?:\.\w+\([^)]*\))*", "", text)
    text = re.sub(r"#h\([^)]*\)", "", text)
    text = re.sub(r"#v\([^)]*\)", "", text)
    text = re.sub(r"#pad\([^]]*\)\[", "", text)
    text = re.sub(r"#text\([^]]*\)\[", "", text)
    text = text.replace(r"\_", "_")
    text = text.replace(r"\%", "%")
    text = text.replace("\\ ", " ")
    text = text.replace("\\", "")
    text = text.replace("~", " ")
    text = text.replace("---", "—")
    text = re.sub(r"(?<!-)--(?!-)", "—", text)
    text = text.replace("→", "—")
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def cleanup_citation_numbers(raw: str) -> str:
    numbers = [part.strip() for part in raw.split(",")]
    return ", ".join(filter(None, numbers))


def sync_content(content_dir: Path, target_dir: Path) -> None:
    target_dir.mkdir(parents=True, exist_ok=True)
    for source in sorted(content_dir.glob("*.typ")):
        target = target_dir / f"{source.stem}.md"
        converted = convert_typst_to_markdown(source.read_text(encoding="utf-8"), source)
        target.write_text(converted.strip() + "\n", encoding="utf-8")


def convert_typst_to_markdown(text: str, source: Path) -> str:
    text = re.sub(r"#context\s+counter\(page\)\.final\(\)\.at\(0\)", "{{NUMPAGES}}", text)
    text = replace_table_macros(text)
    text = replace_code_blocks(text)
    raw_lines = text.splitlines()
    lines: list[str] = []
    in_code_block = False
    i = 0
    while i < len(raw_lines):
        stripped = raw_lines[i].strip()

        if stripped.startswith(CODE_BLOCK_MARKER):
            in_code_block = True
            lines.append(stripped)
            i += 1
            continue
        if stripped == "[[/DOCX-CODE-BLOCK]]":
            in_code_block = False
            lines.append(stripped)
            i += 1
            continue
        if in_code_block:
            lines.append(raw_lines[i].rstrip())
            i += 1
            continue

        if not stripped or stripped.startswith("//") or stripped.startswith("#import"):
            lines.append("")
            i += 1
            continue

        heading_match = re.match(r'#heading\(level:\s*1,\s*numbering:\s*none\)\[(.+)\]', stripped)
        if heading_match:
            lines.append(f"# {cleanup_inline_typst(heading_match.group(1))}")
            lines.append("")
            i += 1
            continue

        fig_match = re.match(r'#fig\("([^"]+)",\s*\[([^\]]+)\](?:.*?width:\s*([^)]+))?\)', stripped)
        if fig_match:
            relative = make_relative_asset_path(fig_match.group(1), source)
            width = normalize_width_spec((fig_match.group(3) or "95%").strip())
            lines.append(f'![{cleanup_inline_typst(fig_match.group(2))}]({relative}){{width={width}}}')
            lines.append("")
            i += 1
            continue

        fig_str_match = re.match(r'#fig\("([^"]+)",\s*"([^"]+)"', stripped)
        if fig_str_match:
            relative = make_relative_asset_path(fig_str_match.group(1), source)
            lines.append(f'![{cleanup_inline_typst(fig_str_match.group(2))}]({relative}){{width=95%}}')
            lines.append("")
            i += 1
            continue

        special_h3 = re.match(r'#pad\([^]]*\)\[#text\([^]]*\)\[([^\]]+)\]\]', stripped)
        if special_h3:
            lines.append(f"{SPECIAL_BOLD_SUBHEADING_MARKER}{cleanup_inline_typst(special_h3.group(1))}")
            lines.append("")
            i += 1
            continue

        if stripped.startswith("Таблица:") or stripped.startswith("|"):
            lines.append(cleanup_table_line(stripped))
            i += 1
            continue

        cleaned = cleanup_inline_typst(stripped)
        if not cleaned:
            lines.append("")
            i += 1
            continue
        if cleaned == "+":
            lines.append("")
            i += 1
            continue

        if cleaned.startswith("= "):
            lines.append(f"# {cleaned[2:].strip()}")
        elif cleaned.startswith("== "):
            lines.append(f"## {cleaned[3:].strip()}")
        elif cleaned.startswith("=== "):
            lines.append(f"### {cleaned[4:].strip()}")
        elif cleaned.startswith("+ "):
            prefix = "1. " if source.name.startswith("06-") else "- "
            lines.append(prefix + cleaned[2:].strip())
        else:
            lines.append(cleaned)
        i += 1

    text_out = "\n".join(collapse_blank_lines(lines)).strip()
    return text_out


def collapse_blank_lines(lines: list[str]) -> list[str]:
    result: list[str] = []
    blank = False
    for line in lines:
        if not line.strip():
            if not blank:
                result.append("")
            blank = True
            continue
        result.append(line)
        blank = False
    return result


def make_relative_asset_path(raw_path: str, source: Path) -> str:
    if raw_path.startswith("/"):
        absolute = ROOT / raw_path.lstrip("/")
    else:
        absolute = (source.parent / raw_path).resolve()
    return Path(relative_path(absolute, CONTENT_DOCX)).as_posix()


def relative_path(path: Path, base: Path) -> str:
    return str(path.relative_to(ROOT)) if path.is_relative_to(ROOT) and base == ROOT else str(Path("..") / path.relative_to(ROOT))


def relative_to_root(path: Path) -> str:
    return path.resolve().relative_to(ROOT.resolve()).as_posix()


def normalize_width_spec(width_raw: str) -> str:
    percent = re.search(r"([\d.]+)%", width_raw)
    if percent:
        return f"{percent.group(1)}%"
    return "95%"


def replace_table_macros(text: str) -> str:
    pieces: list[str] = []
    cursor = 0
    while True:
        match = re.search(r"#(?:long-tbl|tbl)\(", text[cursor:])
        if not match:
            pieces.append(text[cursor:])
            break
        start = cursor + match.start()
        pieces.append(text[cursor:start])
        block = extract_balanced(text, start + match.group(0).index("("), "(", ")")
        full = text[start:start + len(match.group(0)) - 1 + len(block)]
        pieces.append(convert_table_block(full))
        cursor = start + len(full)
    return "".join(pieces)


def convert_table_block(block: str) -> str:
    kind = "long" if block.startswith("#long-tbl") else "short"
    inner = block[block.index("(") + 1 : -1]
    caption_match = re.match(r'\s*"([^"]+)"', inner, re.S)
    caption = caption_match.group(1).strip() if caption_match else "Таблица"
    columns = parse_column_count(inner)
    headers: list[str]

    if kind == "long":
        header_pos = inner.index("header:")
        header_start = inner.index("(", header_pos)
        header_block = extract_balanced(inner, header_start, "(", ")")
        headers = [cleanup_inline_typst(item).strip("*") for item in extract_square_groups(header_block[1:-1])]
        cells = extract_square_groups(inner[header_start + len(header_block) :])
    else:
        cells = extract_square_groups(inner)
        headers = [cleanup_inline_typst(item).strip("*") for item in cells[:columns]]
        cells = cells[columns:]

    rows = [cells[index : index + columns] for index in range(0, len(cells), columns)]
    markdown: list[str] = [f"Таблица: {caption}"]
    markdown.append("| " + " | ".join(escape_md_table(cleanup_inline_typst(cell)) for cell in headers) + " |")
    markdown.append("| " + " | ".join("---" for _ in headers) + " |")
    for row in rows:
        cleaned = [escape_md_table(cleanup_inline_typst(cell)) for cell in row]
        if len(cleaned) < len(headers):
            cleaned.extend([""] * (len(headers) - len(cleaned)))
        markdown.append("| " + " | ".join(cleaned) + " |")
    markdown.append("")
    markdown.append("")
    return "\n".join(markdown)


def parse_column_count(inner: str) -> int:
    match = re.search(r"columns:\s*\(([^)]*)\)", inner)
    if not match:
        return 1
    columns = [part.strip() for part in match.group(1).split(",") if part.strip()]
    return len(columns) or 1


def extract_square_groups(text: str) -> list[str]:
    items: list[str] = []
    depth = 0
    start = None
    for index, char in enumerate(text):
        if char == "[":
            if depth == 0:
                start = index + 1
            depth += 1
        elif char == "]" and depth:
            depth -= 1
            if depth == 0 and start is not None:
                items.append(text[start:index])
                start = None
    return items


def escape_md_table(text: str) -> str:
    return text.replace("|", "\\|")


def is_markdown_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    separator = lines[index + 1].strip()
    return current.startswith("|") and current.endswith("|") and separator.startswith("|") and set(separator.replace("|", "").replace(" ", "")) <= {"-"}


def split_md_row(line: str) -> list[str]:
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return parts


def parse_image_line(line: str) -> dict[str, str | None] | None:
    match = re.match(r'!\[(.*?)\]\((.*?)\)(?:\{width=([^}]+)\})?$', line)
    if not match:
        return None
    return {"caption": match.group(1).strip(), "target": match.group(2).strip(), "width": match.group(3).strip() if match.group(3) else None}


def resolve_asset_path(source_file: Path, target: str) -> Path:
    if target.startswith("/"):
        return ROOT / target.lstrip("/")
    return (source_file.parent / target).resolve()


def width_from_spec(width_spec: str | None, max_width_cm: float) -> float:
    if not width_spec:
        return min(max_width_cm, 14.8)
    percent = re.match(r"([\d.]+)%", width_spec)
    if percent:
        return min(max_width_cm, max_width_cm * float(percent.group(1)) / 100.0 * 0.92)
    absolute = re.match(r"([\d.]+)cm", width_spec)
    if absolute:
        return min(max_width_cm, float(absolute.group(1)) * 0.92)
    return min(max_width_cm, 14.8)


def ensure_figure_caption_suffix(caption: str) -> str:
    normalized = caption.rstrip()
    if normalized.endswith(FIGURE_CAPTION_SUFFIX):
        return normalized
    return f"{normalized}{FIGURE_CAPTION_SUFFIX}"


def build_docx(meta: Meta, content_dir: Path, output_path: Path) -> None:
    builder = DocxBuilder(meta)
    builder.add_title_page()

    abstract = content_dir / "00-abstract.md"
    if abstract.exists():
        builder.add_abstract_file(abstract)

    builder.start_numbered_section()
    builder.add_toc()
    builder.add_page_break()
    builder.add_abbreviations()

    for path in sorted(content_dir.glob("*.md")):
        if path.name in {"00-abstract.md", "README.md"}:
            continue
        builder.add_markdown_file(path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    builder.save(output_path)


def write_content_readme(target_dir: Path) -> None:
    readme = target_dir / "README.md"
    readme.write_text(
        """# content-docx

Здесь хранится Word-ориентированная версия содержимого отчёта.

Структура повторяет папку `content/`, но файлы используют упрощённый Markdown-формат:

- `#`, `##`, `###` для заголовков;
- обычные абзацы;
- списки через `-` и `1.`;
- изображения через `![Подпись](../path/to/image.png){width=95%}`;
- таблицы через обычный Markdown pipe-table, строка `Таблица: ...` перед таблицей используется как подпись.

Файлы можно:

- редактировать вручную под требования Word-версии;
- заново синхронизировать из `content/` командой `python generate_docx.py sync`.
""",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="Генерация DOCX-версии отчёта")
    subparsers = parser.add_subparsers(dest="command", required=True)

    sync_parser = subparsers.add_parser("sync", help="Синхронизировать content-docx из Typst-файлов")
    sync_parser.add_argument("--source", type=Path, default=CONTENT_SRC)
    sync_parser.add_argument("--target", type=Path, default=CONTENT_DOCX)

    build_parser = subparsers.add_parser("build", help="Собрать DOCX из content-docx")
    build_parser.add_argument("--content-dir", type=Path, default=CONTENT_DOCX)
    build_parser.add_argument("--output", type=Path, default=DOCX_OUTPUT)

    args = parser.parse_args()
    meta = parse_meta(META_PATH)

    if args.command == "sync":
        sync_content(args.source, args.target)
        write_content_readme(args.target)
        print(f"Синхронизированы Markdown-файлы в {args.target}")
        return

    if args.command == "build":
        if not args.content_dir.exists():
            raise SystemExit("Папка content-docx не найдена. Сначала выполните sync.")
        build_docx(meta, args.content_dir, args.output)
        print(f"DOCX сохранён: {args.output}")


if __name__ == "__main__":
    main()
