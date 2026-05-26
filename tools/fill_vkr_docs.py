#!/usr/bin/env python3
"""
Заполняет копии шаблонов «Титул ВКР» и «Задание на ВКР» данными студента.

Что делает:
  - меняет все плейсхолдеры (тема, ФИО, шифр, группа, руководитель, цель, задачи, этапы);
  - снимает желтую подсветку (highlight) со всех runs во всех параграфах и таблицах.

Запуск:
  nix-shell -p python3Packages.python-docx --run \
    'python3 tools/fill_vkr_docs.py "Титул Сухов.docx" "Задание Сухов.docx"'
"""
from __future__ import annotations

import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph

STUDENT_FIO = "Сухов Владислав Александрович"
GROUP = "ИКМО-04-24"
SHIFR_PLACEHOLDER = "________"

THEME_LINE_1 = "Управление конфигурацией вычислительной инфраструктуры"
THEME_LINE_2 = "на базе декларативной операционной системы"
THEME_FULL = f"{THEME_LINE_1} {THEME_LINE_2}"

GOAL_TEXT = (
    "разработка системы управления конфигурацией вычислительной инфраструктуры "
    "на основе декларативной операционной системы, обеспечивающей переиспользование "
    "конфигурационных модулей и снижение избыточности конфигурационного кода"
)

TASKS_TEXT = (
    "провести анализ существующих подходов и инструментов управления конфигурацией "
    "вычислительной инфраструктуры; разработать математическую модель оценки качества "
    "управления конфигурацией с системой показателей и весовых коэффициентов; "
    "разработать архитектуру системы с иерархией уровней абстракции (модули, наборы, "
    "профили) и механизмом наследования; сформировать техническое задание в "
    "соответствии с ГОСТ 34.602-2020; реализовать систему и провести экспериментальную "
    "оценку показателей качества"
)

SUPERVISOR_POSITION = "профессор"
SUPERVISOR_FIO = "Краснов Андрей Евгеньевич"

STAGES = [
    (
        "1\n1.1\n1.2\n1.3\n1.4\n1.5",
        (
            "Анализ предметной области и постановка задачи\n"
            "Анализ существующих подходов к управлению конфигурацией\n"
            "Сравнительный анализ моделей управления состоянием инфраструктуры\n"
            "Анализ существующих инструментов управления конфигурацией\n"
            "Постановка цели и задач исследования\n"
            "Разработка математической модели оценки качества"
        ),
        "Исследовательский раздел ВКР, математическая модель",
        "21.03.2026",
    ),
    (
        "2\n2.1\n2.2\n2.3\n2.4\n2.5",
        (
            "Проектирование системы\n"
            "Разработка концептуальной архитектуры системы\n"
            "Проектирование уровня модулей\n"
            "Проектирование уровня наборов модулей\n"
            "Проектирование уровня профилей с механизмом наследования\n"
            "Разработка технического задания по ГОСТ 34.602-2020"
        ),
        "Проектный раздел ВКР, техническое задание",
        "18.04.2026",
    ),
    (
        "3\n3.1\n3.2\n3.3\n3.4\n3.5\n3.6\n3.7",
        (
            "Программная реализация и экспериментальная оценка\n"
            "Выбор и обоснование средств разработки\n"
            "Реализация модулей системы\n"
            "Реализация конфигураций хостов\n"
            "Функциональное тестирование\n"
            "Тестирование развёртывания\n"
            "Сравнительная оценка с аналогами\n"
            "Анализ достижения целевых показателей"
        ),
        "Технологический раздел ВКР, программная реализация",
        "16.05.2026",
    ),
    ("4", "Введение, заключение, список источников, приложения", "Завершённый текст ВКР", "18.05.2026"),
    ("5", "Презентационный материал, аннотация", "Презентация и аннотация", "26.05.2026"),
    ("6", "Нормоконтроль", "Допуск к защите", "27.05.2026"),
]


def replace_in_paragraph(p: Paragraph, old: str, new: str) -> bool:
    """Заменить подстроку в параграфе, корректно работая с runs.

    Объединяет текст всех runs, делает замену, кладёт результат в первый run,
    остальные runs очищает. Форматирование первого run сохраняется.
    """
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not p.runs:
        return False
    p.runs[0].text = new_full
    for r in p.runs[1:]:
        r.text = ""
    return True


def set_paragraph_text(p: Paragraph, new_text: str) -> None:
    """Полностью заменить текст параграфа, сохранив форматирование первого run."""
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def remove_yellow_highlight(doc) -> int:
    """Снять желтую подсветку со всех runs в параграфах и таблицах. Возвращает счётчик."""
    n = 0

    def walk_paragraphs(paragraphs):
        nonlocal n
        for p in paragraphs:
            for r in p.runs:
                if r.font.highlight_color is not None:
                    r.font.highlight_color = None
                    n += 1

    walk_paragraphs(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                walk_paragraphs(cell.paragraphs)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            walk_paragraphs(ncell.paragraphs)
    return n


def set_cell_multiline(cell, lines: str) -> None:
    """Заменить содержимое ячейки набором строк (через перевод строки → отдельные параграфы).

    Сохраняет форматирование первого параграфа ячейки. Лишние параграфы физически
    удаляются (чтобы не оставались пустые «¶ ¶» в конце ячейки), недостающие добавляются.
    Гарантируется минимум один параграф в ячейке.
    """
    parts = lines.split("\n") if lines else [""]
    paragraphs = list(cell.paragraphs)
    while len(paragraphs) < len(parts):
        paragraphs.append(cell.add_paragraph())
    for i, part in enumerate(parts):
        set_paragraph_text(paragraphs[i], part)
    for p in paragraphs[len(parts):]:
        p._element.getparent().remove(p._element)


# =============================================================================
# Титул
# =============================================================================
def fill_title(path: Path) -> None:
    doc = Document(str(path))
    paragraphs = doc.paragraphs

    set_paragraph_text(paragraphs[13], "На тему: " + THEME_LINE_1)
    set_paragraph_text(paragraphs[14], THEME_LINE_2)
    set_paragraph_text(paragraphs[15], "")

    replace_in_paragraph(paragraphs[17], "Дроздова Валерия Александровна", STUDENT_FIO)

    t1 = doc.tables[1]
    t1.rows[0].cells[1].text = SHIFR_PLACEHOLDER
    t1.rows[1].cells[1].text = GROUP

    t2 = doc.tables[2]
    t2.rows[0].cells[2].text = SUPERVISOR_POSITION
    t2.rows[0].cells[3].text = SUPERVISOR_FIO

    cleared = remove_yellow_highlight(doc)
    doc.save(str(path))
    print(f"[Титул] {path.name}: подсветок снято {cleared}")


# =============================================================================
# Задание
# =============================================================================
def fill_assignment(path: Path) -> None:
    doc = Document(str(path))
    paragraphs = doc.paragraphs

    set_paragraph_text(
        paragraphs[7],
        f"1. Тема выпускной квалификационной работы:   {THEME_LINE_1}",
    )
    set_paragraph_text(paragraphs[8], THEME_LINE_2)
    set_paragraph_text(paragraphs[9], "")

    set_paragraph_text(paragraphs[11], f"Цель работы: {GOAL_TEXT}")
    set_paragraph_text(paragraphs[12], f"Задачи работы: {TASKS_TEXT}")

    t2 = doc.tables[2]
    t2.rows[0].cells[1].text = STUDENT_FIO
    t2.rows[0].cells[2].text = STUDENT_FIO
    t2.rows[2].cells[1].text = SHIFR_PLACEHOLDER
    t2.rows[2].cells[2].text = SHIFR_PLACEHOLDER

    t3 = doc.tables[3]
    for i, (num, content, result, deadline) in enumerate(STAGES, start=1):
        row = t3.rows[i]
        set_cell_multiline(row.cells[0], num)
        set_cell_multiline(row.cells[1], content)
        set_cell_multiline(row.cells[2], result)
        set_cell_multiline(row.cells[3], deadline)

    t4 = doc.tables[4]
    t4.rows[1].cells[1].text = SUPERVISOR_POSITION
    t4.rows[1].cells[2].text = SUPERVISOR_FIO

    cleared = remove_yellow_highlight(doc)
    doc.save(str(path))
    print(f"[Задание] {path.name}: подсветок снято {cleared}")


def main() -> None:
    if len(sys.argv) != 3:
        print("usage: fill_vkr_docs.py <title.docx> <assignment.docx>", file=sys.stderr)
        sys.exit(2)
    title_path = Path(sys.argv[1])
    assignment_path = Path(sys.argv[2])
    fill_title(title_path)
    fill_assignment(assignment_path)


if __name__ == "__main__":
    main()
