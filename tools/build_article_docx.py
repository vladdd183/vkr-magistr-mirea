#!/usr/bin/env python3
"""
Генерация научной статьи для РИНЦ в формате .docx по шаблону article/example1.doc.

Параметры оформления (точно по шаблону):

  • Шрифт:              Times New Roman, 14 pt
  • Поля:               левое 30 мм, правое 15 мм, верх/низ 20 мм
  • Межстрочный:        одинарный в «шапке», 1.5 в основном тексте
  • Красная строка:     1.25 см в шапке, 1.5 см в основном тексте
  • Выравнивание:       УДК и авторы — по правому краю
                        Название — по центру (заглавные, жирный)
                        Аннотация и текст — по ширине
  • Подписи таблиц:     над таблицей, обычный шрифт, формат «Таблица 1. …»
  • Подписи рисунков:   под рисунком, жирный, формат «Рисунок 1. …»
  • Источник:           курсив, «Источник: …»
  • Литература:         жирный заголовок, нумерация «1.\\tНазвание»

Запуск:

    nix-shell -p python3Packages.python-docx --run \\
        "python3 tools/build_article_docx.py"

Результат: ``СуховВА_Статья_РИНЦ.docx`` в корне репозитория.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# 📐 Константы оформления (соответствуют XML-параметрам шаблона)
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
SINGLE_SPACING = 1.0           # 240/240 в XML — одинарный (для «шапки»)
ONEHALF_SPACING = 1.5          # 360/240 — полуторный (для основного текста)
DOUBLE_SPACING = 2.0           # 480/240 — двойной (для «Источник:» под таблицей)

INDENT_HEAD = Cm(1.25)         # firstLine="709" twips ≈ 1.25 см (аннотация)
INDENT_BODY = Cm(1.5)          # firstLine="851" twips ≈ 1.50 см (основной текст)

SPACE_BODY_BEFORE = Pt(3)      # before="60" twips ≈ 3 pt
SPACE_BODY_AFTER = Pt(3)       # after="60"  twips ≈ 3 pt


# ---------------------------------------------------------------------------
# 🔧 Вспомогательные функции
# ---------------------------------------------------------------------------


def set_run_font(run, *, size: Pt = FONT_SIZE, bold: bool = False, italic: bool = False) -> None:
    """Применить Times New Roman + размер + (жирный/курсив) к Run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    # Для корректного отображения кириллицы в Word
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_paragraph_format(
    paragraph,
    *,
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: Cm | None = None,
    line_spacing: float = SINGLE_SPACING,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
) -> None:
    """Установить параметры абзаца (выравнивание, отступы, интервалы)."""
    pf = paragraph.paragraph_format
    paragraph.alignment = alignment
    pf.first_line_indent = first_line_indent
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after


def add_para(
    doc,
    text: str = "",
    *,
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: Cm | None = None,
    line_spacing: float = SINGLE_SPACING,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
    bold: bool = False,
    italic: bool = False,
    size: Pt = FONT_SIZE,
):
    """Добавить параграф с одним run и общими параметрами."""
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=alignment,
        first_line_indent=first_line_indent,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
    )
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_runs_para(
    doc,
    runs: list[tuple[str, dict]],
    *,
    alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: Cm | None = None,
    line_spacing: float = SINGLE_SPACING,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
):
    """Добавить параграф с несколькими run'ами (для смешанного форматирования)."""
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=alignment,
        first_line_indent=first_line_indent,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
    )
    for text, opts in runs:
        run = p.add_run(text)
        set_run_font(
            run,
            size=opts.get("size", FONT_SIZE),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    return p


def add_body_para(doc, text: str, **kwargs) -> None:
    """Параграф основного текста: красная строка 1.5 см, полуторный, по ширине."""
    add_para(
        doc,
        text,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_BODY,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
        **kwargs,
    )


def add_section_heading(doc, text: str) -> None:
    """Подзаголовок раздела статьи: красная строка 1.5 см, жирный."""
    add_para(
        doc,
        text,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_BODY,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
        bold=True,
    )


def add_body_runs(doc, runs: list[tuple[str, dict]]) -> None:
    """Основной текст со смешанным форматированием в одном параграфе."""
    add_runs_para(
        doc,
        runs,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_BODY,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
    )


def add_table_caption(doc, text: str) -> None:
    """Подпись таблицы над ней: обычный шрифт, по ширине, красная строка 1.5 см."""
    add_para(
        doc,
        text,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_BODY,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
    )


def add_source_caption(doc, text: str, *, center: bool = False) -> None:
    """Подпись «Источник: …» под таблицей/рисунком, курсив."""
    add_para(
        doc,
        text,
        alignment=(WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY),
        first_line_indent=None,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
        italic=True,
    )


def add_figure_caption(doc, text: str) -> None:
    """Подпись «Рисунок N. …» под рисунком, по центру, жирный."""
    add_para(
        doc,
        text,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
        bold=True,
    )


def set_cell(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    """Записать значение в ячейку таблицы с нужным форматированием."""
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(
        p,
        alignment=(WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT),
        first_line_indent=None,
        line_spacing=SINGLE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0),
    )
    run = p.add_run(text)
    set_run_font(run, size=FONT_SIZE, bold=bold)


def add_data_table(
    doc,
    header: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    """Создать таблицу с границами, заголовком и данными.

    col_widths_cm — ширина колонок в сантиметрах. Если не задан, ширина
    распределяется равномерно. Сумма должна укладываться в рабочую ширину
    листа (16.5 см при полях 30/15 мм на A4).
    """
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _apply_table_borders(table)

    if col_widths_cm is None:
        page_w_cm = 16.5
        per_col = page_w_cm / cols
        col_widths_cm = [per_col] * cols

    # Сначала через tblGrid (общая сетка таблицы)
    tbl = table._element
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gridCol, w in zip(tblGrid.findall(qn("w:gridCol")), col_widths_cm):
            gridCol.set(qn("w:w"), str(int(w * 567)))
            gridCol.set(qn("w:type"), "dxa")
    # Затем через tcW для каждой ячейки (надёжнее в LibreOffice/Word)
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            cell = row.cells[i]
            tcPr = cell._element.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(w * 567)))
            tcW.set(qn("w:type"), "dxa")
    # Запретить автонастройку ширин
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "fixed")

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        set_cell(hdr_cells[i], h, bold=True, center=True)

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            set_cell(cells[c], val, center=(c == 0))

    # Запретить разрыв строк таблицы между страницами
    for row in table.rows:
        trPr = row._element.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
    # Заголовок таблицы повторяется на каждой странице, если таблица
    # переходит на следующую
    hdr_trPr = table.rows[0]._element.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    hdr_trPr.append(tbl_header)


def _apply_table_borders(table) -> None:
    """Установить единичные границы для всех ячеек."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_ref(doc, idx: int, text: str) -> None:
    """Добавить нумерованный источник с табом: «1.\\tНазвание»."""
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
    )
    run = p.add_run(f"{idx}.")
    set_run_font(run)
    tab = OxmlElement("w:tab")
    run._element.append(tab)
    run2 = p.add_run(text)
    set_run_font(run2)


# ---------------------------------------------------------------------------
# 📄 Поля страницы
# ---------------------------------------------------------------------------


def configure_page(doc) -> None:
    """Установить формат A4, поля (30/15/20/20 мм) и базовый шрифт."""
    for section in doc.sections:
        # A4: 210 × 297 мм
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    # Базовый шрифт стиля Normal
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


# ===========================================================================
# 📝 СОДЕРЖАНИЕ СТАТЬИ
# ===========================================================================


UDC = "УДК 004.413:004.451"

AUTHOR_RU = [
    "Сухов Владислав Александрович",
    "Магистрант кафедры математического обеспечения и стандартизации информационных технологий",
    "МИРЭА — Российский технологический университет",
]

AUTHOR_EN = [
    "Sukhov Vladislav Aleksandrovich",
    "Master's student, Department of Mathematical Support and Standardization of Information Technologies",
    "MIREA — Russian Technological University",
]

TITLE_RU = (
    "Модульная архитектура управления конфигурацией вычислительной "
    "инфраструктуры на базе декларативной операционной системы NixOS"
)

TITLE_EN = (
    "A modular configuration management architecture for computing "
    "infrastructure based on the declarative operating system NixOS"
)

ABSTRACT_RU = (
    "В работе рассматривается задача снижения избыточности конфигурационного "
    "кода и устранения configuration drift при управлении вычислительной "
    "инфраструктурой, состоящей из разнородных хостов. Показано, что "
    "традиционные инструменты управления конфигурацией реализуют конвергентную "
    "модель управления состоянием, при которой полное устранение "
    "конфигурационного дрейфа принципиально невозможно. Альтернативой "
    "является конгруэнтная модель, реализованная в декларативных операционных "
    "системах. Предложен интегральный показатель качества управления "
    "конфигурацией, объединяющий коэффициент переиспользования модулей, "
    "степень модульности, сокращение кода и избыточность с обоснованными "
    "весовыми коэффициентами и системой ограничений. Разработана трёхуровневая "
    "архитектура «модули — наборы — профили» с механизмом наследования "
    "профилей, реализующая конгруэнтную модель управления состоянием на "
    "базе декларативной операционной системы NixOS и фреймворка Snowfall Lib. "
    "Экспериментальная оценка реализованной системы, управляющей тремя "
    "хостами различного назначения, продемонстрировала достижение целевых "
    "значений по четырём показателям из шести. Время подготовки нового хоста "
    "сокращено с 60–120 до 10–15 минут. По интегральному показателю качества "
    "разработанная система превосходит распространённые конвергентные "
    "инструменты в 1,6–2 раза. Установлен положительный эффект масштаба: "
    "при расширении инфраструктуры до 10 хостов прогнозируется дальнейший "
    "рост показателей переиспользования и модульности."
)

ABSTRACT_EN = (
    "The paper addresses the problem of configuration code redundancy and "
    "configuration drift in the management of heterogeneous computing "
    "infrastructure. It is shown that traditional configuration management "
    "tools implement a convergent state management model, in which the "
    "complete elimination of configuration drift is in principle unattainable. "
    "An alternative is the congruent model implemented in declarative "
    "operating systems. An integral configuration management quality index "
    "is proposed that combines the module reusability ratio, modularity, "
    "code reduction and redundancy with justified weighting coefficients "
    "and a system of constraints. A three-tier architecture «modules — "
    "suites — profiles» with a profile inheritance mechanism is developed; "
    "the architecture implements a congruent state management model on top "
    "of the declarative operating system NixOS and the Snowfall Lib "
    "framework. Experimental evaluation of the implemented system, "
    "managing three hosts of different purposes, demonstrated the "
    "achievement of the target values for four indicators out of six. "
    "The time required to add a new host was reduced from 60–120 to 10–15 "
    "minutes. With respect to the integral quality index, the proposed "
    "system outperforms common convergent tools by a factor of 1.6 to 2. "
    "A positive scaling effect is established: as the infrastructure "
    "grows to 10 hosts, a further increase in the reusability and "
    "modularity indicators is predicted."
)

KEYWORDS_RU = (
    "управление конфигурацией; Infrastructure as Code; NixOS; "
    "декларативный подход; модульная архитектура; конгруэнтная модель; "
    "Snowfall Lib; DevOps"
)

KEYWORDS_EN = (
    "configuration management; Infrastructure as Code; NixOS; "
    "declarative approach; modular architecture; congruent model; "
    "Snowfall Lib; DevOps"
)


# ===========================================================================
# 🚀 Сборка документа
# ===========================================================================


def build_document(output_path: Path) -> None:
    doc = Document()
    configure_page(doc)

    # Метаданные документа
    core = doc.core_properties
    core.title = TITLE_RU
    core.author = "Сухов Владислав Александрович"
    core.subject = "Научная статья для публикации в РИНЦ"
    core.keywords = KEYWORDS_RU

    # -----------------------------------------------------------------------
    # ШАПКА: УДК, авторы, название, аннотация, ключевые слова
    # -----------------------------------------------------------------------

    # УДК (правый край, жирный, одинарный)
    add_para(
        doc, UDC,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=None,
        line_spacing=SINGLE_SPACING,
        bold=True,
    )

    # Автор RU
    add_para(doc, AUTHOR_RU[0], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING, bold=True)
    add_para(doc, AUTHOR_RU[1], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING)
    add_para(doc, AUTHOR_RU[2], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING)

    # Автор EN
    add_para(doc, AUTHOR_EN[0], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING, bold=True)
    add_para(doc, AUTHOR_EN[1], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING)
    add_para(doc, AUTHOR_EN[2], alignment=WD_ALIGN_PARAGRAPH.RIGHT,
             first_line_indent=None, line_spacing=SINGLE_SPACING)

    # Две пустые строки перед названием
    add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             first_line_indent=None, line_spacing=SINGLE_SPACING)
    add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             first_line_indent=None, line_spacing=SINGLE_SPACING)

    # Название RU/EN (центр, заглавные, жирный)
    add_para(doc, TITLE_RU.upper(), alignment=WD_ALIGN_PARAGRAPH.CENTER,
             first_line_indent=None, line_spacing=SINGLE_SPACING, bold=True)
    add_para(doc, TITLE_EN.upper(), alignment=WD_ALIGN_PARAGRAPH.CENTER,
             first_line_indent=None, line_spacing=SINGLE_SPACING, bold=True)

    # Пустая строка
    add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             first_line_indent=None, line_spacing=SINGLE_SPACING)

    # Аннотация RU (префикс обычный, текст курсив; красная строка 1.25 см; одинарный)
    add_runs_para(
        doc,
        [
            ("Аннотация: ", {"bold": False, "italic": False}),
            (ABSTRACT_RU, {"italic": True}),
        ],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=SINGLE_SPACING,
    )

    # Abstract EN
    add_runs_para(
        doc,
        [
            ("Abstract: ", {"italic": False}),
            (ABSTRACT_EN, {"italic": True}),
        ],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=SINGLE_SPACING,
    )

    # Пустая строка
    add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent=INDENT_HEAD, line_spacing=SINGLE_SPACING)

    # Ключевые слова RU (префикс курсив, текст курсив)
    add_runs_para(
        doc,
        [
            ("Ключевые слова", {"italic": True}),
            (": ", {"italic": False}),
            (KEYWORDS_RU, {"italic": True}),
        ],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=SINGLE_SPACING,
    )

    # Keywords EN
    add_runs_para(
        doc,
        [
            ("Key words: ", {"italic": True}),
            (KEYWORDS_EN, {"italic": True}),
        ],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=SINGLE_SPACING,
    )

    # Пустая строка
    add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent=INDENT_HEAD, line_spacing=SINGLE_SPACING)

    # -----------------------------------------------------------------------
    # ОСНОВНОЙ ТЕКСТ — структура IMRAD «по-русски»
    # -----------------------------------------------------------------------
    # 1) Введение
    # 2) Цель исследования
    # 3) Материал и методы исследования
    # 4) Результаты исследования и их обсуждение
    # 5) Выводы
    # -----------------------------------------------------------------------

    # --- Введение ---
    add_section_heading(doc, "Введение.")
    add_body_para(
        doc,
        "В условиях цифровой трансформации сложность вычислительной инфраструктуры "
        "организаций непрерывно возрастает. Каждая среда — от рабочих станций "
        "разработчиков до серверов и кластерных узлов — требует согласованной "
        "настройки десятков и сотен параметров операционной системы и прикладных "
        "сервисов. Согласно отчёту Uptime Institute «Annual Outage Analysis 2023», "
        "45 % респондентов указали ошибки конфигурации и управления изменениями "
        "в качестве одной из наиболее распространённых причин сетевых отказов, "
        "а для крупных простоев на уровне ИТ-систем доля подобных инцидентов "
        "составила 40 % [1]. По данным New Relic, медианная годовая стоимость "
        "ИТ-простоев достигает 7,75 млн долл., а 32 % организаций оценивают "
        "стоимость критических простоев более чем в 500 тыс. долл. в час [2]."
    )

    add_body_para(
        doc,
        "Дополнительной системной проблемой является конфигурационный дрейф "
        "(configuration drift) — постепенное расхождение фактического состояния "
        "хостов с документированным описанием. Подобный дрейф формируется "
        "вследствие ручных правок, несогласованных процессов развёртывания и "
        "временных исправлений, которые не фиксируются в системах управления "
        "версиями [3]. Совместное исследование EMA и Itential показало, что "
        "только 34 % специалистов по сетевому администрированию полностью "
        "удовлетворены существующими инструментами управления конфигурациями, "
        "а 75 % организаций обеспокоены возможными ошибками при их применении [4]."
    )

    add_body_para(
        doc,
        "Традиционные инструменты управления конфигурацией (Ansible, Puppet, "
        "Chef) реализуют конвергентную модель управления состоянием: декларация "
        "описывает целевое состояние, после чего агент последовательно сравнивает "
        "текущее состояние с целевым и применяет корректирующие действия. "
        "Подобная модель в принципе не исключает configuration drift, поскольку "
        "описание целевого состояния практически всегда неполно и в системе "
        "сохраняются побочные эффекты предыдущих операций [5]. Альтернативой "
        "является конгруэнтная модель, реализованная в декларативных операционных "
        "системах NixOS и GNU Guix: состояние системы полностью определяется "
        "декларативным описанием и атомарно пересоздаётся при каждом изменении [6, 7]."
    )

    add_body_para(
        doc,
        "Принцип DRY (Don't Repeat Yourself) [8] широко применяется в разработке "
        "прикладного программного обеспечения, однако в области управления "
        "конфигурацией инфраструктуры методики его систематического применения "
        "проработаны недостаточно. При прямолинейном использовании NixOS для "
        "нескольких хостов возникает избыточность: одинаковые настройки "
        "копируются из одной конфигурации в другую, что нивелирует ключевое "
        "преимущество декларативного подхода."
    )

    # --- Цель исследования ---
    add_section_heading(doc, "Цель исследования.")
    add_body_para(
        doc,
        "Целью исследования является разработка системы управления конфигурацией "
        "вычислительной инфраструктуры на базе декларативной операционной системы "
        "NixOS, обеспечивающей максимизацию переиспользования конфигурационных "
        "модулей и минимизацию избыточности кода, а также количественная оценка "
        "достигаемого эффекта."
    )
    add_body_para(
        doc,
        "Для достижения цели в работе решены следующие задачи: разработана "
        "математическая модель оценки качества управления конфигурацией; "
        "спроектирована трёхуровневая модульная архитектура (модули — наборы — "
        "профили с механизмом наследования); реализована система для управления "
        "тремя хостами различного назначения; проведена экспериментальная оценка "
        "показателей качества и сравнение с конвергентными аналогами."
    )

    # --- Материал и методы исследования ---
    add_section_heading(doc, "Материал и методы исследования.")
    add_body_para(
        doc,
        "В исследовании применяется классификация моделей управления состоянием, "
        "выделяющая три типа: дивергентную, конвергентную и конгруэнтную. При "
        "дивергентном управлении состояние системы изменяется неконтролируемо "
        "посредством ручных интервенций, что неизбежно приводит к появлению "
        "уникальных «хостов-снежинок». Конвергентное управление сравнивает "
        "текущее состояние с целевым и последовательно приближает их друг к "
        "другу, не гарантируя при этом исключение configuration drift [5]. "
        "Конгруэнтное управление обеспечивает принудительное соответствие "
        "системы её декларативному описанию — состояние полностью пересоздаётся "
        "при каждом изменении конфигурации [7]. Сравнение моделей по ключевым "
        "характеристикам приведено в таблице 1."
    )

    # Таблица 1: модели управления состоянием
    add_table_caption(doc, "Таблица 1. Сравнение моделей управления состоянием инфраструктуры")
    add_data_table(
        doc,
        ["Характеристика", "Дивергентная", "Конвергентная", "Конгруэнтная"],
        [
            ["Предсказуемость", "Низкая", "Средняя", "Высокая"],
            ["Воспроизводимость", "Отсутствует", "Частичная", "Полная"],
            ["Configuration drift", "Неизбежен", "Возможен", "Исключён"],
            ["Откат изменений", "Невозможен", "Сложен", "Тривиален"],
            ["Примеры", "Ручное SSH", "Ansible, Puppet", "NixOS, Guix"],
        ],
        col_widths_cm=[5.0, 3.8, 3.9, 3.8],
    )
    add_source_caption(doc, "Источник: составлено автором")

    add_body_para(
        doc,
        "Для количественного сравнения архитектурных решений в работе предложен "
        "интегральный показатель качества управления конфигурацией Q_CM, который "
        "рассчитывается как линейная комбинация четырёх метрик: Q_CM = ω₁·R + "
        "ω₂·M + ω₃·S − ω₄·C → max, где R — коэффициент переиспользования модулей "
        "(reusability), M — степень модульности архитектуры (modularity), S — "
        "коэффициент сокращения кода (size reduction), C — избыточность "
        "конфигурации (code redundancy, штрафной фактор), ω₁,…,ω₄ — весовые "
        "коэффициенты, удовлетворяющие условию Σ ω_i = 1."
    )

    add_body_para(
        doc,
        "Факторы вычисляются по следующим формулам. Коэффициент переиспользования "
        "модулей оценивает долю кода, повторно использованного более чем одним "
        "хостом, по отношению к суммарному объёму конфигурационного кода с учётом "
        "кратности применения: R = Σ LOC(m)·(U(m) − 1) / Σ LOC(m)·U(m), где M — "
        "множество модулей, LOC(m) — объём кода модуля m в строках, U(m) — число "
        "хостов, использующих модуль. Степень модульности равна доле модулей, "
        "применяемых не менее чем двумя хостами: M = |{m ∈ M : U(m) ≥ 2}| / |M|. "
        "Коэффициент сокращения кода характеризует относительное уменьшение объёма "
        "конфигурации хоста при переходе от изолированного описания к модульному: "
        "S = 1 − LOC_modular / LOC_isolated. Избыточность C оценивается как доля "
        "дублирующихся фрагментов кода (3-строчных k-грамм) в общем объёме "
        "конфигурации: C = LOC_dup / LOC_total."
    )

    add_body_para(
        doc,
        "Весовые коэффициенты установлены экспертно с учётом приоритетов, "
        "характерных для систем управления конфигурацией: ω₁ = 0,35 "
        "(переиспользование), ω₂ = 0,25 (модульность), ω₃ = 0,25 "
        "(сокращение кода), ω₄ = 0,15 (штраф за избыточность). Система "
        "ограничений для целевых значений показателей имеет вид: R ≥ 0,50; "
        "M ≥ 0,60; S ≥ 0,40; C ≤ 0,10; T_deploy ≤ 15 минут. При достижении "
        "минимальных целевых значений Q_CM = 0,35·0,50 + 0,25·0,60 + 0,25·0,40 "
        "− 0,15·0,10 = 0,41, что и принимается в качестве порогового значения "
        "интегрального показателя."
    )

    add_body_para(
        doc,
        "Для устранения избыточности конфигурационного кода при сохранении "
        "конгруэнтной модели управления предложена трёхуровневая архитектура, "
        "включающая уровни модулей, наборов и профилей. Модули представляют "
        "собой атомарные, переиспользуемые конфигурационные компоненты, "
        "отвечающие за один аспект системы (например, сервис Docker, "
        "аудиоподсистема, настройки SSH). Каждый модуль предоставляет "
        "параметризованный интерфейс через атрибут options и применяет "
        "конфигурацию по флагу enable. Наборы (suites) — логически связанные "
        "группы модулей, объединённых по функциональному назначению. Профили "
        "определяют роли машин и поддерживают наследование через атрибут "
        "extends; наследование является аддитивным."
    )

    add_body_para(
        doc,
        "В качестве операционной системы выбрана NixOS 25.11 — декларативный "
        "дистрибутив Linux, реализующий конгруэнтную модель управления через "
        "пакетный менеджер Nix [7]. Для структурирования Nix Flakes применён "
        "фреймворк Snowfall Lib, обеспечивающий стандартизированную структуру "
        "директорий и автоматическое обнаружение модулей и хостов [9]. "
        "Удалённое развёртывание реализовано через утилиту deploy-rs с "
        "механизмом Magic Rollback: при отсутствии подтверждения успешного "
        "применения конфигурации в течение 60 секунд хост автоматически "
        "возвращается к предыдущему поколению."
    )

    add_body_para(
        doc,
        "Управляющий репозиторий vladOS-v2 содержит 25 NixOS-модулей общим "
        "объёмом 2 397 строк кода, 7 модулей home-manager (659 строк), "
        "5 наборов (570 строк) и 13 профилей (6 desktop- и 7 server-профилей). "
        "Общий объём кодовой базы составляет около 6 200 строк в более чем "
        "70 Nix-файлах. Численные значения показателей R, M, S, C и Q_CM "
        "получены автоматически по исходному коду репозитория с помощью "
        "разработанного скрипта на Python."
    )

    # --- Результаты исследования и их обсуждение ---
    add_section_heading(doc, "Результаты исследования и их обсуждение.")
    add_body_para(
        doc,
        "Разработанная система используется для централизованного управления "
        "тремя физическими хостами различного назначения: рабочей станцией "
        "разработчика (профиль developer, 221 строка специфичной конфигурации), "
        "рабочим узлом Kubernetes-кластера с GPU NVIDIA (профиль "
        "server-k3s-agent, 150 строк) и узлом, совмещающим функции "
        "CI-раннера и файлового хранилища (профили server-ci и server-storage, "
        "205 строк). В среднем объём специфичного для хоста кода составил "
        "192 строки против 500 строк при изолированной конфигурации, что "
        "соответствует сокращению на 62 %."
    )

    add_body_para(
        doc,
        "Результаты автоматического расчёта показателей качества разработанной "
        "системы и их сопоставление с целевыми значениями представлены в "
        "таблице 2."
    )

    # Таблица 2: показатели качества
    add_table_caption(doc, "Таблица 2. Численные значения показателей качества разработанной системы")
    add_data_table(
        doc,
        ["Показатель", "Формула", "Значение", "Цель"],
        [
            ["R — переиспользование", "(801·2 + 666) / 4665", "0,49", "≥ 0,50"],
            ["M — модульность", "13 / 25", "0,52", "≥ 0,60"],
            ["S — сокращение кода", "1 − 576 / 1500", "0,62", "≥ 0,40"],
            ["C — избыточность", "140 / 1875", "0,07", "≤ 0,10"],
            ["Q_CM — индекс качества", "линейная свёртка", "0,44", "≥ 0,41"],
            ["T_deploy — время развёртывания", "—", "10–15 мин", "≤ 15 мин"],
        ],
        col_widths_cm=[5.5, 5.2, 2.8, 3.0],
    )
    add_source_caption(doc, "Источник: расчёт автора по исходному коду системы vladOS-v2")

    add_body_para(
        doc,
        "Полученное значение Q_CM = 0,44 превышает целевое пороговое значение "
        "0,41. Из шести нормативных показателей четыре достигнуты полностью "
        "(S, C, Q_CM, T_deploy); коэффициент переиспользования R достигнут на "
        "98 %, а модульность M — на 87 %. Недостижение полных целевых значений "
        "по R и M объясняется малым числом управляемых хостов (3): при "
        "увеличении количества хостов ранее реализованные модули автоматически "
        "используются повторно, что приводит к улучшению обоих показателей."
    )

    add_body_para(
        doc,
        "Для оценки эффекта от расширения инфраструктуры выполнен сценарный "
        "прогноз показателей при увеличении числа хостов до 10. Прогноз "
        "построен исходя из предположения, что новые хосты соответствуют "
        "типовым ролям, для которых в системе уже определены профили "
        "(developer, server-base, server-k3s-agent, server-ci, server-storage, "
        "server-virt). Результаты прогнозного расчёта приведены в таблице 3."
    )

    # Таблица 3: прогноз масштабирования
    add_table_caption(doc, "Таблица 3. Прогноз показателей качества при масштабировании инфраструктуры")
    add_data_table(
        doc,
        ["Показатель", "3 хоста", "10 хостов (прогноз)", "Изменение"],
        [
            ["R", "0,49", "0,65", "+33 %"],
            ["M", "0,52", "0,88", "+69 %"],
            ["S", "0,62", "0,60", "−3 %"],
            ["C", "0,07", "0,04", "−43 %"],
            ["Q_CM", "0,44", "0,59", "+34 %"],
        ],
        col_widths_cm=[3.0, 3.5, 5.5, 4.5],
    )
    add_source_caption(doc, "Источник: расчёт автора (сценарный прогноз)")

    add_body_para(
        doc,
        "Существенный рост R и M при умеренном изменении S свидетельствует о "
        "положительном эффекте масштаба: разработанная архитектура становится "
        "тем эффективнее, чем больше хостов управляется из единого репозитория."
    )

    add_body_para(
        doc,
        "В рамках функционального тестирования выполнены 10 тестовых случаев, "
        "охватывающих сборку конфигураций всех трёх хостов, применение базовых "
        "и наследуемых профилей, включение типовых модулей (Docker, K3s), "
        "активацию набора common и генерацию визуализации топологии. Все "
        "тестовые случаи пройдены успешно, что подтверждает корректность "
        "работы механизма композиции и наследования. В рамках тестирования "
        "развёртывания проверены пять сценариев: локальное применение "
        "конфигурации занимает 45 с, удалённое развёртывание через deploy-rs — "
        "от 2 мин 15 с до 3 мин 40 с в зависимости от объёма передаваемых "
        "данных. Механизм Magic Rollback корректно срабатывает при искусственно "
        "смоделированном сбое активации, возвращая хост в предыдущее поколение "
        "через 60 с. Ручной откат через nixos-rebuild --rollback выполняется "
        "за 12 с."
    )

    add_body_para(
        doc,
        "Введение системы профилей с наследованием существенно сокращает "
        "трудозатраты на интеграцию нового хоста в инфраструктуру. При ручной "
        "настройке изолированной NixOS-конфигурации подготовка нового хоста "
        "занимает 60–120 минут в зависимости от роли. В разработанной системе "
        "аналогичная процедура сводится к созданию директории хоста, генерации "
        "hardware-configuration.nix стандартной утилитой nixos-generate-config "
        "и написанию 10–15 строк, описывающих применяемый профиль и специфичные "
        "параметры. Полное время подготовки составляет 10–15 минут, что отвечает "
        "поставленному ограничению T_deploy ≤ 15 минут."
    )

    add_body_para(
        doc,
        "Для оценки разработанной системы относительно существующих решений "
        "произведено сравнение по введённой системе показателей с тремя "
        "эталонными подходами: Ansible (императивный/конвергентный, наиболее "
        "распространённый инструмент управления конфигурацией), Puppet "
        "(декларативный конвергентный) и изолированными NixOS-конфигурациями "
        "(декларативный конгруэнтный, но без модульной архитектуры). Результаты "
        "сравнения приведены в таблице 4."
    )

    # Таблица 4: сравнение с аналогами
    add_table_caption(doc, "Таблица 4. Сравнение количественных показателей с аналогами")
    add_data_table(
        doc,
        ["Показатель", "Ansible", "Puppet", "Изол. NixOS", "vladOS (3 х.)", "vladOS (10 х.)"],
        [
            ["R", "0,30", "0,35", "0,00", "0,49", "0,65"],
            ["M", "0,40", "0,45", "0,10", "0,52", "0,88"],
            ["S", "0,20", "0,25", "0,00", "0,62", "0,60"],
            ["C", "0,25", "0,20", "0,40", "0,07", "0,04"],
            ["Q_CM", "0,22", "0,27", "−0,04", "0,44", "0,59"],
        ],
        col_widths_cm=[3.0, 2.0, 2.0, 2.8, 3.2, 3.5],
    )
    add_source_caption(doc, "Источник: расчёт автора, значения для Ansible/Puppet получены экспертным анализом типовых репозиториев")

    add_body_para(
        doc,
        "Полученное значение Q_CM = 0,44 превосходит Ansible в 2 раза (0,22), "
        "Puppet — в 1,6 раза (0,27) и недостижимо для изолированной "
        "NixOS-конфигурации (отрицательное значение, обусловленное "
        "избыточностью). Преимущество разработанной системы достигается "
        "главным образом за счёт высокого коэффициента сокращения кода "
        "S = 0,62 (вдвое выше, чем у конвергентных аналогов) и низкой "
        "избыточности C = 0,07 (почти в четыре раза ниже, чем у Ansible). "
        "Помимо количественных метрик, разработанная система превосходит "
        "аналоги по ряду качественных характеристик: конгруэнтная модель "
        "управления исключает configuration drift по построению, атомарные "
        "обновления гарантируют отсутствие частично применённых состояний, "
        "а возможность мгновенного отката к любому из ранее применённых "
        "поколений снижает риск при изменении конфигурации до уровня, "
        "недостижимого в конвергентных системах."
    )

    add_body_para(
        doc,
        "Полученные результаты согласуются с современными тенденциями в "
        "области управления конфигурацией. Согласно отчёту Puppet «State of "
        "DevOps Report 2024», 80 % организаций практикуют DevOps, причём "
        "50 % из них относятся к категории высокопроизводительных [10]. "
        "Рынок DevOps-инструментов прогнозируется на уровне 25,5 млрд долл. "
        "к 2028 году [11], а 45 % организаций уже внедрили практику "
        "Infrastructure as Code, при этом 74 % ИТ-руководителей рассматривают "
        "её в качестве ключевого фактора облачной трансформации [12]. "
        "Разработанная система реализует все ключевые принципы Infrastructure "
        "as Code (декларативность, воспроизводимость, версионирование) с "
        "дополнительной гарантией конгруэнтности, что отвечает запросам "
        "современной отрасли на повышение надёжности и предсказуемости "
        "управления инфраструктурой."
    )

    # --- Выводы ---
    add_section_heading(doc, "Выводы.")
    add_body_para(
        doc,
        "В работе предложена и реализована модульная система управления "
        "конфигурацией вычислительной инфраструктуры на базе декларативной "
        "операционной системы NixOS. Разработана математическая модель оценки "
        "качества управления конфигурацией с интегральным показателем Q_CM, "
        "учитывающим переиспользование модулей, модульность архитектуры, "
        "сокращение кода и избыточность. Предложена трёхуровневая архитектура "
        "«модули — наборы — профили» с механизмом наследования профилей, "
        "обеспечивающим аддитивное накопление функциональности."
    )

    add_body_para(
        doc,
        "Экспериментальная оценка реализованной системы на инфраструктуре из "
        "трёх хостов показала достижение интегрального показателя качества "
        "Q_CM = 0,44 при пороговом значении 0,41, что соответствует превышению "
        "порога на 7 %. Коэффициент сокращения кода составил 62 % (в 2–3 раза "
        "выше, чем у конвергентных аналогов), избыточность — 7 % (в 3–6 раз "
        "ниже). Время подготовки нового хоста сокращено с 60–120 минут до "
        "10–15 минут. По интегральному показателю Q_CM разработанная система "
        "превосходит Ansible в 2 раза и Puppet в 1,6 раза. Установлено, что "
        "разработанная архитектура демонстрирует положительный эффект масштаба: "
        "при увеличении количества управляемых хостов до 10 прогнозное значение "
        "Q_CM возрастает до 0,59, что свидетельствует о пригодности предложенного "
        "подхода для управления гетерогенными инфраструктурами среднего и "
        "крупного масштаба."
    )

    add_body_para(
        doc,
        "Дальнейшие направления исследования включают расширение прогнозной "
        "модели на инфраструктуры с числом хостов более 50, разработку методики "
        "автоматизированной генерации профилей по описанию ролей, интеграцию "
        "с системами GitOps и адаптацию предложенной архитектуры к управлению "
        "контейнерными конфигурациями (Docker, Kubernetes-манифесты) на базе "
        "Nix-выражений."
    )

    # Пустая строка перед литературой
    add_body_para(doc, "")

    # -----------------------------------------------------------------------
    # ЛИТЕРАТУРА
    # -----------------------------------------------------------------------

    # «Литература:» — жирный заголовок с красной строкой 1.5 см
    p_lit = doc.add_paragraph()
    set_paragraph_format(
        p_lit,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=INDENT_HEAD,
        line_spacing=ONEHALF_SPACING,
        space_before=SPACE_BODY_BEFORE,
        space_after=SPACE_BODY_AFTER,
    )
    r = p_lit.add_run("Литература:")
    set_run_font(r, bold=True)

    references = [
        "Annual Outage Analysis 2023: The Causes and Impacts of IT and Data "
        "Center Outages [Электронный ресурс] / Uptime Institute. – 2023. – "
        "URL: https://uptimeinstitute.com/uptime_assets/"
        "5f40588be8d57272f91e4526dc8f821521950b7bec7148f815b6612651d5a9b3-"
        "annual-outages-analysis-2023.pdf (дата обращения: 14.05.2026).",

        "Industry's Largest Survey Finds Enterprises Realize 2X ROI on "
        "Observability [Электронный ресурс] / New Relic. – 12.09.2023. – "
        "URL: https://newrelic.com/press-release/20230912 "
        "(дата обращения: 14.05.2026).",

        "Configuration Drift: Why It's Bad and How to Eliminate It "
        "[Электронный ресурс] / Aqua Security. – 2022. – "
        "URL: https://www.aquasec.com/cloud-native-academy/"
        "vulnerability-management/configuration-drift/ "
        "(дата обращения: 14.05.2026).",

        "The State of Network Automation: Configuration Management Obstacles "
        "are Universal [Электронный ресурс] / EMA Research, Itential. – 2021. – "
        "URL: https://www.itential.com/what-is-network-automation/"
        "research-state-of-network-automation/ (дата обращения: 14.05.2026).",

        "Declarative Management of Kubernetes Objects Using Configuration "
        "Files [Электронный ресурс] / Kubernetes Documentation. – 2023. – "
        "URL: https://kubernetes.io/docs/tasks/manage-kubernetes-objects/"
        "declarative-config/ (дата обращения: 14.05.2026).",

        "AWS Well-Architected Framework. REL08-BP04: Deploy using immutable "
        "infrastructure [Электронный ресурс] / Amazon Web Services. – "
        "10.04.2023. – URL: https://docs.aws.amazon.com/wellarchitected/"
        "2023-04-10/framework/rel_tracking_change_management_"
        "immutable_infrastructure.html (дата обращения: 14.05.2026).",

        "NixOS Manual [Электронный ресурс] / NixOS Project. – "
        "URL: https://nixos.org/manual/nixos/stable/ "
        "(дата обращения: 14.05.2026).",

        "What is the DRY (don't repeat yourself) principle? "
        "[Электронный ресурс] / TechTarget. – 04.08.2025. – "
        "URL: https://www.techtarget.com/whatis/definition/DRY-principle "
        "(дата обращения: 14.05.2026).",

        "Snowfall Lib Documentation [Электронный ресурс] / Snowfall Software. – "
        "URL: https://snowfall.org/guides/lib/quickstart/ "
        "(дата обращения: 14.05.2026).",

        "State of DevOps Report 2024 [Электронный ресурс] / Puppet by Perforce. "
        "– 2024. – URL: https://www.puppet.com/resources/state-of-devops-report "
        "(дата обращения: 14.05.2026).",

        "DevOps Market Size, Share & Trends Analysis Report "
        "[Электронный ресурс] / Polaris Market Research. – 2024. – "
        "URL: https://www.polarismarketresearch.com/industry-analysis/"
        "devops-market (дата обращения: 14.05.2026).",

        "Infrastructure as Code (IaC) Market Size, Share & Segmentation "
        "(Global Forecast 2024–2032) [Электронный ресурс] / SNS Insider. – "
        "2024. – URL: https://www.snsinsider.com/reports/"
        "infrastructure-as-code-market-4659 (дата обращения: 14.05.2026).",

        "ГОСТ 34.602-2020. Информационные технологии. Комплекс стандартов на "
        "автоматизированные системы. Техническое задание на создание "
        "автоматизированной системы. – М.: Стандартинформ, 2020. – 27 с.",

        "ISO/IEC 25010:2011. Systems and Software Engineering. Systems and "
        "Software Quality Requirements and Evaluation (SQuaRE). System and "
        "Software Quality Models. – Geneva: ISO, 2011. – 34 p.",

        "ANSI/EIA-649B. National Consensus Standard for Configuration "
        "Management. – Warrendale: SAE International, 2011. – 48 p.",
    ]

    for i, ref in enumerate(references, start=1):
        add_ref(doc, i, ref)

    # Сохраняем
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Сохранено: {output_path}")


# ---------------------------------------------------------------------------
# 🚀 Точка входа
# ---------------------------------------------------------------------------


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    else:
        out = Path(__file__).resolve().parent.parent / "СуховВА_Статья_РИНЦ.docx"
    build_document(out)
