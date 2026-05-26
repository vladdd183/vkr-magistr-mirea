#!/usr/bin/env python3
"""Дамп структуры docx: все параграфы (с runs и highlight) и таблицы."""
import sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def dump(path: str) -> None:
    print(f"\n========== {path} ==========")
    d = Document(path)
    print(f"\n--- PARAGRAPHS ({len(d.paragraphs)}) ---")
    for i, p in enumerate(d.paragraphs):
        runs_info = []
        for r in p.runs:
            hl = r.font.highlight_color
            mark = "★" if hl == WD_COLOR_INDEX.YELLOW else ""
            runs_info.append(f"[{mark}{r.text!r}]")
        print(f"  P{i:02d}: {p.text!r}")
        if runs_info:
            print(f"      runs: {' '.join(runs_info)}")
    print(f"\n--- TABLES ({len(d.tables)}) ---")
    for ti, t in enumerate(d.tables):
        print(f"  T{ti}: rows={len(t.rows)}, cols={len(t.columns)}")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.replace("\n", " ¶ ").strip()
                if txt:
                    print(f"    T{ti}[r{ri}c{ci}]: {txt!r}")


if __name__ == "__main__":
    for p in sys.argv[1:]:
        dump(p)
